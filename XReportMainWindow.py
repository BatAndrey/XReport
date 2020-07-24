# -*- coding: utf-8 -*-

import sys
import docx
from docx.document import Document
import time
import configparser
import os
from PyQt5 import QtWidgets, QtCore
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QComboBox, QTextEdit, QPushButton, QAction, QMainWindow, \
    QLineEdit, QApplication, QLayout, QVBoxLayout, QHBoxLayout, QWidget, QFrame, QFormLayout, QStatusBar, QProgressDialog
from PyQt5.QtCore import (QTimer)
import XReportSettings
import XReportSaveMenu

dir_templ = r'C:\Users\a.batischev\Desktop\templ\templ2.docx'

class XReportMainWindow(QMainWindow):
    tmpl_dir_path = ''
    check_frequence = 2000

    def __init__(self):
        super().__init__()
        self.read_config(file_name='config.ini')
        self.initGui()
        self.exchange_templates_combo()
        # self.set_timer_connect_email()
        self.ok_btn_compile_report()

    def read_config(self, file_name):
        config = configparser.ConfigParser()
        config.read(file_name)
        self.tmpl_dir_path = config['FILE_DIR_PATH']['dir_path']
        print(self.tmpl_dir_path)

    def initGui(self):
        self.main_window = QWidget()
        self.timer_connect_email = QTimer()

        self.setGeometry(400, 400, 700, 400)
        self.setWindowTitle('X-Report')

        #self.doc = docx.Document(dir_templ)

        self.templ_text_edid = QTextEdit()
        self.templ_combo_box = QComboBox()
        self.btn_connect = QPushButton("connect email")
        self.btn_compile_report = QPushButton("компилировать отчет")
        self.inicialis = QLineEdit()
        self.age = QLineEdit()
        self.data = QLineEdit()
        self.doctor = QLineEdit()

        self.vbox_text_combo = QVBoxLayout()
        self.vbox_text_combo.addWidget(self.templ_text_edid)
        self.vbox_text_combo.addWidget(self.templ_combo_box)
        self.vbox_text_combo.addWidget(self.btn_connect)

        self.vbox_data_person = QVBoxLayout()

        self.form_person = QFormLayout()
        self.form_person.addRow("ФИО", self.inicialis)
        self.form_person.addRow("Возраст", self.age)
        self.form_person.addRow("Дата", self.data)
        self.form_person.addRow("Врач", self.doctor)
        self.form_person.addRow(self.btn_compile_report)
        self.form_person.addRow(self.vbox_data_person)

        self.data.setInputMask("Дата приема: 99.B9.9999; _")
        self.age.setInputMask("Дата рождения: 99.B9.9999; _")

        self.doctor.setText("Ксения Константиновна")


        self.hbox_all = QHBoxLayout()
        self.hbox_all.addLayout(self.vbox_text_combo)
        self.hbox_all.addLayout(self.form_person)

        self.main_window.setLayout(self.hbox_all)
        self.setCentralWidget(self.main_window)

        self.file_menu = self.menuBar().addMenu('File')

        self.show_settings_action = QAction('Settings')
        self.show_save_action = QAction('Save')

        self.show_settings_action.triggered.connect(self.show_settings)
        self.file_menu.addAction(self.show_settings_action)
        #self.show_save_action.triggered.connect(self.show_save)
        self.show_save_action.triggered.connect(self.save_to_file)
        self.file_menu.addAction(self.show_save_action)

        # self.progress_dialog = QProgressDialog("Operation in progress.", "Cancel", 0, 100)
        # self.progress_dialog.show()
        # self.progress_dialog.startTimer(10)

    def exchange_templates_combo(self):
        item_for_combo = ['ogk', 'bone', 'air']
        self.templ_combo_box.addItems(item_for_combo)
        self.templ_combo_box.currentTextChanged.connect(self.set_text_templates)

    def set_text_templates(self):
        file_name = self.templ_combo_box.currentText()
        templ_text = self.read_file_templ(file_name=file_name)
        self.templ_text_edid.setText(templ_text)

        self.templ_text_edid.textChanged.connect(lambda: self.templ_combo_box.setDisabled(
            self.templ_text_edid.toPlainText() != ''))

    def read_file_templ(self, file_name):
        file_for_read = open(self.tmpl_dir_path + file_name + '.txt', 'r')
        tmp_text = file_for_read.readlines()
        file_for_read.close()
        return ' '.join(tmp_text)

    # def set_timer_connect_email(self):
    #     self.timer_connect_email.timeout.connect(self.check_email)
    #     self.timer_connect_email.start(self.check_frequence)

    # def check_email(self):
    #     print('mail exist')
    #     self.statusBar().showMessage('Checking mail...', 2000)
    #     # connection = ConnectionEmail()
    #     # if connection.exist_unread_msg():
    #     #     #self.statusBar().showMessage('Unread message exist')
    #     #     self.statusBar().messageChanged('Unread message exist')
    #     #     connection.download_msg()
    #     #     connection.close()
    #     # else:
    #     #     print('Unread message unexist')
    #     #     connection.close()

    # def ok_btn_connect_email(self):
    #     self.btn_connect.clicked.connect(self.check_email)

    def show_settings(self):
        self.settings_widget = XReportSettings.SettingsWidget()
        self.settings_widget.show()
        self.settings_widget.setAttribute(Qt.WA_DeleteOnClose)
        self.settings_widget.settings_changed.connect(self.apply_settings)

    def apply_settings(self, freq, dir_path):
        self.tmpl_dir_path = dir_path
        self.check_frequency = freq
        print('apply settings')
        self.timer_connect_email.start(self.check_frequency)

    def show_save(self):
        self.save_wedget = XReportSaveMenu.SaveMenu()
        self.save_wedget.show()
        self.save_wedget.setAttribute(Qt.WA_DeleteOnClose)
        self.save_wedget.settings_changed.connect(self.apply_save)

    def apply_save(self):
        pass

    def save_to_file(self):
        options = QtWidgets.QFileDialog()
        self.fileName, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save To File", r'C:\Users\a.batischev\Desktop\templ\report_redy', "All (*);;Text Files (*.txt)", options=options)
        if self.fileName:
            self.writeFile = open(self.fileName, 'w', encoding='utf-8')
            self.writeFile.write(self.templ_text_edid.toPlainText())
            self.writeFile.close()
            self.statusbar.showMessage('Saved to %s' % self.fileName)


    def ok_btn_compile_report(self):
        print('press compile btn')
        self.btn_compile_report.pressed.connect(self.compil_report)

    def compil_report(self):
        # print('work slots')
        # get_date = self.data.text()
        # #print('date', get_date)
        # get_inicial = self.inicialis.text()
        # #print('inicials', get_inicial)
        # get_doctor = self.doctor.text()
        # get_text = self.templ_text_edid.toPlainText()

        self.doc = docx.Document(dir_templ)

        self.doc.add_paragraph(self.data.text() + '\n')
        self.doc.add_paragraph(self.inicialis.text() + '\n')
        self.doc.add_paragraph(self.age.text() + '\n')
        self.doc.add_paragraph(self.templ_text_edid.toPlainText() + '\n')
        self.doc.add_paragraph(self.doctor.text())

        self.doc.save('temple22.docx')





if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = XReportMainWindow()
    ex.show()
    # unittest.main()
    sys.exit(app.exec_())
