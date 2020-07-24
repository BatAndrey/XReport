import docx
from docx.document import Document
dir_templ = r'C:\Users\a.batischev\Desktop\templ\templ.docx'

doc = docx.Document(dir_templ)
text = []
for paragraph in doc.paragraphs:
    text.append(paragraph.text)
print('\n'.join(text))


p1 = doc.add_paragraph('hello*3')
doc.save('templ_ex.docx')
doc2 = docx.Document(dir_templ)
docx.document.Document.add_paragraph(doc2, 'da')
docx.document.Document.save(doc2, path_or_stream=dir_templ)

Document.add_paragraph(doc2, 'hop hei lalalay')
doc2.save('templ_ex2.docx')