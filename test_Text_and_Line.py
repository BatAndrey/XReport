from PyQt5.QtWidgets import QMainWindow, QTextEdit, QLineEdit, QLabel, QLayout, QHBoxLayout, \
    QAction, QFormLayout, QPushButton, QWidget, QVBoxLayout, QApplication
import sys
import docx
from docx.document import Document
dir_templ = r'C:\Users\a.batischev\Desktop\templ\templ.docx'
class Test_Text_and_Line(QMainWindow):


    def __init__(self):
        super().__init__()
        self.initGUI()
        self.signal_text_exchange()
        #self.create_templ()

    def initGUI(self):
        self.main_window = QWidget()

        self.setGeometry(400, 400, 700, 400)
        self.setWindowTitle('X-Report')

        self.templ_text_edid = QTextEdit()

        self.btn_compile_report = QPushButton("компилировать отчет")

        self.doc = docx.Document(dir_templ)

        self.inicialis = QLineEdit()
        self.age = QLineEdit()
        self.data = QLineEdit()
        self.doctor = QLineEdit()

        self.hbox = QHBoxLayout()
        self.hbox.addWidget(self.templ_text_edid)

        self.vbox_data_person = QVBoxLayout()

        self.form_person = QFormLayout()
        self.form_person.addRow("ФИО", self.inicialis)
        self.form_person.addRow("Возраст", self.age)
        self.form_person.addRow("Дата", self.data)
        self.form_person.addRow("Врач", self.doctor)
        self.form_person.addRow(self.btn_compile_report)
        self.form_person.addRow(self.vbox_data_person)

        self.data.setInputMask("Дата приема: 99.B9.9999; _")
        self.age.setInputMask("Дата рождения: 99.B9.9999; _")

        self.doctor.setText("Ксения Константиновна")

        self.hbox_all = QHBoxLayout()
        self.hbox_all.addLayout(self.hbox)
        self.hbox_all.addLayout(self.form_person)

        self.main_window.setLayout(self.hbox_all)
        self.setCentralWidget(self.main_window)



    def compil_report(self):
        dir_templ = r'C:\Users\a.batischev\Desktop\templ\templ.docx'
        get_date = self.data.text()
        get_inicial = self.inicialis.text()
        get_doctor = self.doctor.text()
        get_text = self.templ_text_edid.toPlainText()
        get_age = self.age.text()
        print(get_text, get_inicial, get_date, get_age, get_doctor)


    def create_templ(self):
        self.doc.add_paragraph(self.data.text())
        self.doc.add_paragraph(self.inicialis.text())
        self.doc.add_paragraph(self.age.text())
        self.doc.add_paragraph(self.templ_text_edid.toPlainText())
        self.doc.add_paragraph(self.doctor.text())

        self.doc.save('temple22.docx')


        #self.templ_text_edid.setText(get_text)

    def signal_text_exchange(self):
        self.btn_compile_report.clicked.connect(self.create_templ)
        # self.templ_text_edid.textChanged.connect(self.compil_report)
        # self.inicialis.textChanged.connect(self.compil_report)
        # self.data.textChanged.connect(self.compil_report)
        # self.age.textChanged.connect(self.compil_report)
        # self.doctor.textChanged.connect(self.compil_report)

        # self.data.textChanged.connect(self.line_to_text)
        # self.inicialis.textChanged.connect(self.line_to_text)
        # # self.age.textChanged.connect(self.line_to_text)
        # self.doctor.textChanged.connect(self.line_to_text)

    def line_to_text(self):
        self.templ_text_edid.setText(self.data.text())
        self.templ_text_edid.setText(self.inicialis.text())
        # self.templ_text_edid.setText(self.age.text())
        # self.templ_text_edid.setText(self.doctor.text())



if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = Test_Text_and_Line()
    ex.show()
    # unittest.main()
    sys.exit(app.exec_())